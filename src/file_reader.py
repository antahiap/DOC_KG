from PyPDF2 import PdfReader
import json
import glob
import os
import re
import pandas as pd
import requests
import shutil
from pathlib import Path
import numpy as np
import docx
from constants import PAPER_IMG_PATH, PAPER_TXT_PATH, PAPER_REF_PATH

from dotenv import load_dotenv
load_dotenv()

class FileReader():
    def __init__(self, 
                 file_path, 
                 write_json=True,
                 write_text=False,
                 update=False
                 ):
        self.missing_tags = []

        if not file_path =='':
            self.file_path = file_path
            self.pdf_file_name = file_path.stem
        self.write_json = write_json
        self.write_text = write_text
        self.update = update

    def batch_read_pdf(self, src_dir):
        pdf_list = glob.glob(os.path.join(src_dir, '*.pdf'))

        for self.file_path in pdf_list:
            # input(self.file_path)
            pdf_file_name = os.path.basename(self.file_path)
            self.pdf_file_name, _ = os.path.splitext(pdf_file_name)

            self.pdf_reader = PdfReader(open(self.file_path, 'rb'))
            self.num_pages = len(self.pdf_reader.pages)

            self._pdf_img()
            self._pdf_to_txt()
    
    def read_pdf(self, dl_ref=None):

        file_extension = self.file_path.suffix   

        if file_extension == '.pdf':

            pdf_file = open(self.file_path, 'rb')
            self.dl_ref = dl_ref
            self.pdf_reader = PdfReader(pdf_file)
            self.num_pages = len(self.pdf_reader.pages)

            # self._pdf_img()
            txt_dic = self._pdf_to_txt()
            self._write_content(txt_dic)
            self.content = txt_dic
            return txt_dic
        
        elif file_extension == '.docx':
            doco = Doco(self.file_path)
            self.missing_tags = doco.missing_tags
            content = doco.get_contents()
            self._write_content(content)

            self.content = content
            return content

        else:
            print(f'Unsupported file format: {file_extension}')
    
    def get_text_path(self):
        return PAPER_TXT_PATH / f"{self.pdf_file_name}.txt"
    
    def get_json_path(self):
        return PAPER_TXT_PATH / f"{self.pdf_file_name}.json"
    
    def get_json(self):
        json_path = self.get_json_path()
        with open(json_path, 'r') as f:
            data = json.load(f)
        return data

    def _pdf_img(self):
        for page_number in range(self.num_pages):
            page = self.pdf_reader.pages[page_number]
            count = 0
            for image_file_object in page.images:
                img_path = PAPER_IMG_PATH / f'{self.pdf_file_name}_{count}_{image_file_object.name}'
                with open(img_path,  "wb") as fp:
                    fp.write(image_file_object.data)
                    count += 1
                
    def _split_sections(self, ref_list):

        def _overlap_matches(pattern, text):
            offset = 0
            matches = []

            while True:
                # Search for matches starting from the current offset
                match = re.search(pattern, text[offset:])

                if not match:
                    break
                
                # Get the matched substring and add it to the list of matches
                matches.append(match.group(1))

                # Update the offset to search for the next match
                offset += match.start() + 1
            return(matches)

        def _get_text(matches, text):
            text_sec = []
            for i in range(len(matches)):
                start = matches_1[i]

                if i==len(matches_1)-1:
                    end = 'Reference'
                else:
                    end = matches_1[i+1]

                text_i = text.split(start)[1].split(end)
                text_i = text_i[0]
                text_sec.append(text_i.strip())

            return(text_sec)
        
        def _sort_sections(matches_1):    
            missing = []
            pdf_strcture = []    
            for mi, match in enumerate(matches_1):
                # if text_sec[mi] == '':
                #     print(match, self.pdf_file_name)
                #     id, section_name, pos = '', match, [0]
                #     continue
                if mi == 0:
                    section_number = ''
                    section_name = match
                    pos = [0]
                    id = ''

                else:
                    try:
                        match_split = re.findall(pattern_2, match, re.DOTALL)[0]
                        section_number = match_split[0].strip('.').split('.')
                        section_name = match_split[1]
                        id = match_split[0]

                        pos = [None for x in section_number]
                        for j in range(len(section_number)):
                            pos[j] = int(section_number[j]) -1

                    except IndexError:
                        section_name = match
                        try:
                            text_sec[mi] = re.findall(
                                r'([A-Za-z0-9 \n\.\-\,\;\(\)]+)', text_sec[mi])[0]
                        except:
                            text_sec[mi] = text_sec[mi]
                        pos = [0]
                        id = ''


                if section_name[1] == ' ':
                    section_name = section_name[:1] + section_name[2:]
                section_name =  section_name[0].upper() + section_name[1:].lower()
                json_tmplt = {
                        'id': id,
                        'section': section_name,
                        'text': text_sec[mi],
                        'subsection': []
                    }

                try:
                    if len(pos) == 1:
                        pdf_strcture.append(json_tmplt)
                    elif len(pos) == 2:
                        pdf_strcture[pos[0]]['subsection'].append(json_tmplt)  
                    elif len(pos) == 3:
                        pdf_strcture[pos[0]]['subsection'][pos[1]]['subsection'].append(json_tmplt)  
                    elif len(pos) == 4:
                        pdf_strcture[pos[0]]['subsection'][pos[1]]['subsection'].append(json_tmplt)  
                    else:
                        missing.append(json_tmplt)

                except IndexError:

                    # the text has line with starting a number
                    print(pos)
                    missing.append(match)
                    print(match)
                    continue

            return pdf_strcture, missing

        # Get Abstract
        pattern_0 = r'(Abstract)\n'
        abstract = re.findall(pattern_0, self.text, flags= re.DOTALL | re.IGNORECASE)

        exclude_words = ['Figure', 'Table', 'fig']
        
        # match section headings with the numaber
        pattern_1 =  r'[^' + '|'.join(map(re.escape, exclude_words)) + r']\n(\d\.?\d?\.?\d?\.? [A-Z].+)\n' 
        matches_1 = abstract + _overlap_matches(pattern_1, self.text) 
        text_sec = _get_text(matches_1, self.text)   # text of each section
        
        # section headings splited numaber
        pattern_2 =  r'(\d\.?\d?\.?\d?\.?) ([A-Z][\-A-Za-z *]+)'    


        if len(matches_1) < 4: 
            # paper doesn't match to our regex
            missing = []
            pdf_strcture = [{'id':'', 'section': 'Full Paper', 'text':self.text}]
        else:
            pdf_strcture, missing = _sort_sections(matches_1)


        out_data = pdf_strcture + [{'missing': missing}]
        out_data = out_data + [{'references': ref_list}]
        out_data = out_data + [{
            'title': self.title, 
            'arxiv_id': self.pdf_file_name
            }]
        
        return out_data

    def _write_content(self, content_dic):
        txt_path = self.get_text_path()
        if self.write_text:
            with open(txt_path, "w", encoding="utf-8") as output_file:
                output_file.write(self.text)
        
        json_path = self.get_json_path()
        if self.write_json:
            with open(json_path, 'w', encoding='utf-8') as json_file:
                json.dump(content_dic, json_file, indent=4) 

    def _pdf_to_txt(self):
        text = ""
        TITLE =[]
        FNT_SIZE = []

        def visitor_body(text, cm, tm, fontDict, fontSize):
            y = tm[5]
            x = tm[4]

            if y > 40 and y < 800:
                if x > 35 and x <1000:
                    parts.append(text)

                    if fontSize >14 and fontSize<17:
                        TITLE.append(text) #.strip())
                        FNT_SIZE.append(fontSize)
        
        def clean_title():
            title = arxiv_title()
            if title:
                return title
            
            max_fnt = np.median(FNT_SIZE)
            title = [TITLE[i] for i, ti in enumerate(FNT_SIZE) if ti>max_fnt]
            return title
        
        def arxiv_title():
            import requests 

            arxiv_id = self.pdf_file_name
            api_url = f'https://export.arxiv.org/api/query?id_list={arxiv_id}'

            # Make a GET request to the arXiv API
            response = requests.get(api_url)

            # Check if the request was successful
            if response.status_code == 200:
                # Parse the response to extract the title
                xml_data = response.text
                title_start = xml_data.find("<title>") + len("<title>")
                title_end = xml_data.find("</title>", title_start)
                title = xml_data[title_start:title_end]
                return(title)
            else:
                return None
        
        json_path = self.get_json_path()
        if os.path.isfile(json_path):
            if not self.update:
                return self.get_json()
        
        for page_num in range(self.num_pages):
            parts = []

            page = self.pdf_reader.pages[page_num]  #pdf_reader.getPage(page_num)
            page.extract_text(
                visitor_text=visitor_body)

            # remove page number
            if not parts[-1].endswith('\n'): parts= parts[:-1]
            text += "".join(parts)
        
        self.text = text
        self.title=''.join(clean_title())
        ref_list = self._get_refernces()
        content_dic = self._split_sections(ref_list)


        return(content_dic)

    def _get_refernces(self):

        def _clean_ref_id(ref_list):
            ref_ids = []
            for ref in ref_list:
                id = ref.split('[')[1].split(']')[0]
                if id in ref_ids:
                    ref_ids.remove(id)
                ref_ids.append(id)

            return ref_ids
        
        def _get_ref_text(ids):
            old_pattern = 'Reference.*\n'
            ref_txt = []
            
            for i, id in enumerate(ids):

                end_pattern = f'\[{id}\](.*?\.)\n|$'

                try:
                    id_next = f'\[{ids[i+1]}\]'
                    ref_pattern = old_pattern + f'\[{id}\](.*)\n{id_next}'

                except IndexError:
                    ref_pattern = old_pattern + end_pattern

                text = re.findall(ref_pattern, self.text, flags=re.DOTALL | re.IGNORECASE)

                # Needed for two column papers
                if not text == []:
                    text = text[0]
                    old_pattern = old_pattern + f'\[{id}\].*\n'
                else:
                    ref_pattern = old_pattern + end_pattern
                    text = re.findall(ref_pattern, self.text, flags=re.DOTALL)[0]
                    old_pattern = ''

                text = ' '.join(text.split('\n'))
                ref_txt.append(text)

            return ref_txt

        ref_pattern = r'(\[\d+[a-zA-Z]?\].*)'
        references = re.findall(ref_pattern, self.text)
        ref_ids = _clean_ref_id(references)
        ref_texts = _get_ref_text(ref_ids)

        return ref_texts
    
    def list_reference(self, ref_texts):

        def _arxiv_dl(paper):

            arxiv_pattern = r'arXiv:(\d+\.\d+)'
            arxiv_id = re.findall(arxiv_pattern, paper, flags=re.IGNORECASE)
            try:
                arxiv_id = arxiv_id[0]
            except IndexError:
                return []
            
            ref_path = PAPER_REF_PATH / Path(f"{arxiv_id}.pdf")
            if os.path.isfile(ref_path):
                return([arxiv_id])

            arxiv_url = f"https://export.arxiv.org/pdf/{arxiv_id}"
            response = requests.get(arxiv_url)

            if response.status_code == 200:

                with open(ref_path, 'wb') as pdf_file:
                    pdf_file.write(response.content)

                print(f"PDF downloaded and saved as {ref_path}")
                return([arxiv_id])
            else:
                print("Failed to download the PDF.")
                return([])

        def _google_dl(paper):
            from serpapi import GoogleSearch
            api_key = os.getenv("SERP_API_KEY")

            search_query = f'"{paper}" filetype:pdf'
            search = GoogleSearch({'q': search_query, 'api_key': api_key})
 
 
            results = search.get_dict()
 
            pdf_links = []

            for result in results.get('organic_results', []):
                link = result.get('link', '')
                match = re.search(r'(http[s]?://[^\s]+\.pdf)', link)
                if match:
                    pdf_links.append(match.group())
            # 
            return
           
        pdf_in_ref = PAPER_REF_PATH / Path(self.pdf_file_name +'.pdf')
        if not os.path.isfile(pdf_in_ref):
            shutil.copy(self.file_path, pdf_in_ref)
        
        ref_paper = []
        for paper in ref_texts:

            # check for arxiv paper
            download = _arxiv_dl(paper)
            if len(download) == 0:
                continue
                # download = _google_dl(paper)
            else:
                ref_paper += download
                
        return(ref_paper)


class Doco(object):
    def __init__(self, doc_path):

        self.doc_name_orgnl = doc_path.name

        self.doc_name = self._normalize_name(doc_path.name)
        self.path = doc_path.parent

        self.doc = docx.Document(doc_path)
        self.missing_tags = []
        # self.g = Graph()
        # self.doc_uri = URIRef(f"http://example.org/{doc_name}")

        # self.doco = Namespace("http://purl.org/spar/doco/")
        # self.g.bind("doco", self.doco)

        return 

    def _extract_titles(self):
        titles = []
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):  # Heading styles indicate titles/subtitles
                
                self.g.add((self.doc_uri, self.doco.SectionTitle, para_text))
        return titles

    def _normalize_name(self, name):
        # Normalize the name by removing leading/trailing spaces and converting to lowercase and remove white spaces 
        return name.strip().lower().replace(" ", "_")
    
    def extract_table(self):
        """Extracts tables with their content in a structured format, removing empty rows."""
        tables = []

        for table_idx, table in enumerate(self.doc.tables, start=1):
            table_data = {
                "table_number": table_idx,
                "rows": []
            }
            for row in table.rows:
                # Check if the row is empty (if all cells in the row are empty)
                row_data = [cell.text.strip() for cell in row.cells]
                if any(cell for cell in row_data):  # Only include the row if any cell contains text
                    table_data["rows"].append(row_data)

            # Only add the table if it has at least one row with data
            if table_data["rows"]:
                tables.append(table_data)

        return tables
    def _get_paragraphs(self):

        """Extracts paragraphs with numbering and text content."""

        para_counter = 0

        data = []
        for i, para in enumerate(self.doc.paragraphs):

            para_text = para.text.strip()
            tag = para.style.name
            
            if para_text == '':
                continue

            if (tag.startswith('Heading') or tag.startswith('Appendix') or tag == 'Title'
                or tag.startswith('Ebene')):

                if tag.startswith('Heading'):
                    tag_split = 'Heading '
                elif tag.startswith('Appendix'):
                    tag_split = 'Appendix '
                elif tag.startswith('Ebene'):
                    tag_split = 'Ebene_'
                else: tag_split = None

                if tag_split:
                    id = tag.split(tag_split)[-1].replace(" ", "")
                else: id = 1


                data.append({
                    "id": id,
                    "section": para_text,
                    "text": '',
                    "subsection": [],
                    "tag": 'section',
                })

                # self.g.add((self.doc_uri, self.doco.SectionTitle, para_text)) 

            elif tag in ['Caption', 'Abbindung', 'TableBody', 'Caption1']:
                data.append({
                    "id": "0",
                    "section": '',
                    "text": para_text,
                    "subsection": [],
                    "tag": 'caption',
                })

            elif tag.startswith("toc") or tag in [
                'Absatzebene1', 'table of figures']:

                # Table of contents

                data.append({
                    "id": '0',
                    "section": '',
                    "text": para_text,
                    "subsection": [],
                    "tag": 'table_of_contents',
                })

            elif tag == 'Standard+3 Punkt':
                # skip list 
                pass
            
            # elif tag =='Default':
            #     print(para_text)
            #     print('--------------------------------')

            else:
                if not tag in [
                'Normal','List Paragraph', 'Body', 'No Spacing', 'annotation text',
                'VSn_7_Fett und unterstrichen', 'Normal (Web)', 'Body Text', '0030',
                'normal1',  'Intense Quote'
                ]:
                    if not tag in self.missing_tags: 
                        self.missing_tags.append(tag)
                
                data.append({
                    "id": "0",
                    "section": '',
                    "text": para_text,
                    "subsection": [],
                    "tag": 'paragraph',
                })
                


            para_counter += 1

        return data
    
    def _restructure(self, data):

        restructed_data = []
        i = 0
        id_hold =0

        while i < len(data):

            section = data[i]['section']
            text = data[i]['text']

            if int(data[i]['id']) >= id_hold or int(data[i]['id'])  == 0:
                id_hold =  int(data[i]['id']) 
                i+=1

                if i < len(data): 
                    while int(data[i]['id']) == 0:
                        texti = data[i]['text']
                        text += texti + "\n"
                        i+= 1

                        if i == len(data):
                            break
                                               
                new_item = { 
                    "id": str(id_hold) ,
                    "section": section,
                    "text": text,
                    "subsection": [],
                    "tag": "section"
                }
                i-=1
                id_hold = int(data[i]['id'])
            

            else:
                new_item = data[i]
                id_hold = int(data[i]['id'])

            i+=1

            restructed_data.append(new_item)


        return restructed_data

    def element_id(self, i, tag):
        return f"{self.doc_name}_{tag}{i+1}"
    
    def get_contents(self):

        data = self._get_paragraphs()
        data = self._restructure(data)

        #data.append(self.extract_table())
        data.append({
            "title": "",
            "arxiv_id": self.doc_name,
            "text": self.doc_name,
            "tag": "document",
            "section": ''
        })
        self.content = data
        return data


def get_files_in_directory(path, ext):
    pdf_pattern = os.path.join(path, '**', f'*.{ext}')  
    pdf_files = glob.glob(pdf_pattern, recursive=True)

    return pdf_files

if __name__ == '__main__':
    file_src =FileReader(Path("data/article_pdf/1603.06147.pdf")) #2308.16441.pdf")) #1706.03762.pdf")) 1308.0850.
    # issue: 2309.03409

    all_files = get_files_in_directory('../data', 'docx')
    missing_labels = []
    jj = 1

    for j in range(jj, len(all_files)):
        f = all_files[j]
        print(j, f)
        f_path = Path(f)

        file_src =FileReader(f_path)
        text_dic = file_src.read_pdf()
        missing_labels += file_src.missing_tags


    print(missing_labels)


